from __future__ import annotations

import sys

import docx
import os
import json
import report
from post_message import send_message
import re
from collections import Counter
from typing import cast
from jira import JIRA
from jira.client import ResultList
from jira.resources import Issue
import requests
from datetime import datetime, timedelta, date
from verification import verific
from init import init_path

# ---------------------------------------------- функция форматирования даты ------------------------------

def check_structure_ip(name_ip):
    with open('check.json', 'r', encoding='utf-8') as f: # открытие файла для чтения
        data = json.load(f) # чтение данных из файла в формате JSON в объект Python

    if name_ip in data:
        structure = data[name_ip]
    else:
        structure = 'нет совпадения'
    return structure

def formating_date (stroka):
    #day = stroka[1:3]
    day = stroka[1:stroka.index(' ')-1]
    if len(day) != 2:
        if day == '1': day = '01'
        if day == '2': day = '02'
        if day == '3': day = '03'
        if day == '4': day = '04'
        if day == '5': day = '05'
        if day == '6': day = '06'
        if day == '7': day = '07'
        if day == '8': day = '08'
        if day == '9': day = '09'
    #print('День', day)
    buf_month = stroka[stroka.index(' ')+1:stroka.index(' 20')]
    #print('Месяц', buf_month)
    year = stroka[-7:-3]
    #print('Год', year)

    month = ''
    if buf_month == 'января': month = '01'
    elif buf_month == 'февраля': month = '02'
    elif buf_month == 'марта': month = '03'
    elif buf_month == 'апреля': month = '04'
    elif buf_month == 'мая': month = '05'
    elif buf_month == 'июня': month = '06'
    elif buf_month == 'июля': month = '07'
    elif buf_month == 'августа': month = '08'
    elif buf_month == 'сентября': month = '09'
    elif buf_month == 'октября': month = '10'
    elif buf_month == 'ноября': month = '11'
    elif buf_month == 'декабря': month = '12'
    else: month = '-'
    if month != '-':                                            # если месяц найден и год равен текущему году
        date_act_form = year + '-' + month + '-' + day
        global date_for_rename
        date_for_rename = day + '-' + month + '-' + year                    # переменная для переименования акта
        if month == '': date_act_form = "Проверьте корректность даты акта"
        date_act_form = datetime.strptime(date_act_form, "%Y-%m-%d").date()  #перевод в тип даты
        return date_act_form
    else:
        date_for_rename = '-'  # переменная для переименования акта
        return -1

# -----------------------------------------------поиск и запись всех путей файлов------------------------

paths = []
folder = os.path.dirname(init_path())  # получение пути из функции инициализации конфигурационного файла
for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))
#print('Всего актов', len(paths))
#print(paths)
#----------------------завершение работы при отсутствии актов----------------------------
if len(paths) == 0:
    print('в дирректории нет ни одного акта для проверки')
    input("Нажмите enter для завершения...")
    sys.exit()

# -----------------------------------------------создание экземпляра документа------------------------
for path in paths:
    print(f'>>> Получение данных из акта {path}...')
    doc = docx.Document(path)
    properties = doc.core_properties
    """print('Автор документа:', properties.author)
    print('Автор последней правки:', properties.last_modified_by)
    print('Дата создания документа:', properties.created)
    print('Дата последней правки:', properties.modified)
    print('Дата последней печати:', properties.last_printed)
    print('Количество сохранений:', properties.revision)"""

# -----------------------------------------------работа с таблицами------------------------
    mas_tables=[]
    for table in doc.tables:
        mas_tables.append(table)
        """for row in table.rows:
            for cell in row.cells:
                print(cell.text)"""

    #print(len(mas_tables))

# -----------------------------------------------инициализация переменных------------------------
    number_act = 0
    number_zayavka = 0
    date_act = ''
    project_act = ''
    key_project_act = ''
    period = ''
    time_act = 0
    rate_act = 0
    rate_zayavka = 0
    project_cost_act = 0
    total_cost_act = 0
    total_cost_zayavka = 0
    name_act2 = ''
    name_act3 = ''
    name_user = ''
    # -----------------------------------------------получение данных------------------------

    #number_act = mas_tables[0].cell(0, 0).text[-1]  # Номер акта
    #number_zayavka = mas_tables[4].cell(0, 0).text[-1] # Номер заявки

    #date_act = mas_tables[1].cell(0, 1).text  # Дата акта
    #date_act_f=formating_date(date_act)       #форматированная дата
    #print('Дата акта -', date_act_f)

    project_act = mas_tables[0].cell(1, 1).text  # Наименование проекта
    #print(project_act)
    project_act = project_act[project_act.index('ПРОЕКТ: ')+8:len(project_act)+1]
    project_act = project_act.replace('\n', ' ')
    project_act = project_act.replace('ё', 'е')
    #key_project_act = mas_tables[2].cell(1, 2).text  # Ключ проекта

    period = mas_tables[0].cell(1, 3).text  # период
    date_start = period[2:12].replace(' ', '')  # дата начала
    date_end = period[15:26].replace(' ', '')  # дата завершения

    time_act = mas_tables[0].cell(1, 4).text  # трудозатраты
    time_act = time_act.replace(',', '.')
    time_act = time_act[time_act.index('Отработано:') + 11:len(time_act) + 1]

    rate_act = mas_tables[0].cell(1, 5).text  # ставка в акте
    if rate_act.find('Ставка:') != -1:
        rate_act = rate_act[rate_act.index('Ставка:') + 7:len(rate_act) + 1]
    else:
        rate_act = mas_tables[0].cell(1, 5).text
    #print(rate_act)
    rate_zayavka = mas_tables[2].cell(1, 1).text  # ставка в заявке

    project_cost_act = mas_tables[0].cell(1, 6).text  # стоимость по проекту
    project_cost_act = project_cost_act[project_cost_act.index('Стоимость:') + 10:len(project_cost_act) + 1]

    total_cost_act = mas_tables[0].cell(2, 6).text  # Итого в акте
    total_cost_act = total_cost_act[total_cost_act.index('ИТОГО:') + 6:len(total_cost_act) + 1]

    total_cost_zayavka = mas_tables[2].cell(1, 3).text  # Итого в заявке

    rate_for_calculate = rate_act.replace(',', '.')
    cost_for_verification = format(float(rate_for_calculate) * float(time_act), '.2f')  #расчет стоимости проекта

    #name_act2 = mas_tables[3].cell(1, 1).text[3:len(mas_tables[3].cell(1, 1).text)]  # имя в  акте 2
    #name_act3 = mas_tables[6].cell(1, 1).text[3:len(mas_tables[3].cell(1, 1).text)]  # имя  в акте 3


    #print('Номер акта -', number_act)
    #print('Номер заявки -', number_zayavka)
    #print(date_act)  # Дата акта
    #print('Наименование проекта -', project_act)  # Наименование проекта
    #print('Ключ проекта -', key_project_act)  # Ключ проекта
    #print('Период проверки -', period)  # период
    #print('Начало периода', date_start)
    #print('Конец периода', date_end)
    #print('Трудозатраты в акте -', time_act)  # трудозатраты
    #print('Ставка в акте -', rate_act)  # ставка в Акте
    #print('Ставка в заявке -', rate_zayavka)  # ставка в заявке
    #print('Стоимость по проекту -', project_cost_act)  # стоимость по проекту
    #print('Итого в таблице -', total_cost_act)  # Итого
    #print('Итого в заявке -', total_cost_zayavka)
    #print('Имя в подписи 1 -', name_act2)
    #print('Имя в подписи 2 -', name_act3)

# -----------------------------------------------работа с текстом------------------------
    text = []
    for paragraph in doc.paragraphs:                   # получение списка параграфов
        text.append(paragraph.text)

    #print('\n'.join(text))

    # ----------------------------------------------- выделение имени из акта ------------------------
    start_index = text[4].find('ИП ')
    type_of_act = 0                      # тип акта 0 - ИП, 1 - Самозанятый
    if start_index != -1:
        full_name_act1 = text[4][text[4].index('ИП ')+3:text[4].index(', именуемый')]   # выделение имени
        name_act2 = mas_tables[1].cell(1, 1).text[3:len(mas_tables[1].cell(1, 1).text)]  # имя в  акте 2
        name_act3 = mas_tables[3].cell(1, 1).text[3:len(mas_tables[3].cell(1, 1).text)]  # имя  в акте 3
        type_of_act = 0
    if start_index == -1:
        full_name_act1 = text[4][text[4].index('стороны, и ')+11:text[4].index(', именуемый')]   # выделение имени
        name_act2 = mas_tables[1].cell(1, 1).text[0:len(mas_tables[1].cell(1, 1).text)]  # имя в  акте 2
        name_act3 = mas_tables[3].cell(1, 1).text[0:len(mas_tables[3].cell(1, 1).text)]  # имя  в акте 3
        type_of_act = 1
    #print(full_name_act1)

    name_act = full_name_act1[0:full_name_act1.rfind(' ')]
    name_act = name_act.replace(' ', '')
    name_act = name_act.replace('ё', 'е')
    #print('Имя в акте -', name_act)

    # -----------------------------------------------выделение суммы из текста акта------------------------

    #print(text[8])
    rub = text[8][text[8].index('сумму ')+6:text[8].index(' (')]   # выделение суммы руб
    copeyka = text[8][text[8].index(' копе')-2:text[8].index(' копе')]   # выделение суммы копейки
    #print(rub)
    #print(copeyka)

    total_cost_act_in_text = rub + ',' + copeyka   # полная сумма
    #print('Итого в тексте -', total_cost_act_in_text)

#----------------------------------------------------- выделение из текста ---------------------

    #number_act = mas_tables[0].cell(0, 0).text[-1]  # Номер акта
    number_act = text[0][text[0].index('АКТ  № ') + 7:len(text[0]) + 1]
    #print('Номер акта -', number_act)

    #number_zayavka = mas_tables[4].cell(0, 0).text[-1] # Номер заявки
    number_zayavka = text[19][text[19].index('Заявка на оказание услуг №') + 26:len(text[19]) + 1]
    #print('Номер заявки -', number_zayavka)
# ---------------------------------- парсинг даты ---------------------------------------
    #date_act = mas_tables[1].cell(0, 1).text  # Дата акта
    index1 = 0
    index_count = 1

    for i in range(len(text[2])):
        if i > 50 and not text[2][i].isspace():
            index1 = index_count
            break
        index_count +=1

    date_act = text[2][index1-1:len(text[2])]
    date_act_f=formating_date(date_act)       #форматированная дата
    # print('Дата акта -', date_act_f)
    print('OK')
# --------------------------------------- Запуск модуля jira---------------------------------
    print('>>> Подключение к JIRA...')
    jira = JIRA(server='https://jira.i-sol.eu', basic_auth=('tcontrol', '1J*iBJGT'))

    # ------------------------ Ввод ФИО---------------------
    date1 = date_start
    date2 = date_end
    date1 = date1.replace('.', '/')
    date2 = date2.replace('.', '/')
    date1 = datetime.strptime(date1, "%d/%m/%Y")
    date2 = datetime.strptime(date2, "%d/%m/%Y")

    datedelta = date.today() - timedelta(33)            # установка дельты для даты в - 35 дней для сокращения выгрузки данных из джира
    datedelta = datedelta.strftime("%Y-%m-%d")
    # print(date1)
    # print(date2)
    # name_user = input('Введите Фамилию и Имя: \n')

    print('OK')
#--------------------------------- проверка есть внутри ИП другие участники -----------------------
    print('>>> Выгрузка дынных из JIRA...')
    new_name = check_structure_ip(full_name_act1)
    if new_name == 'нет совпадения':  # проверяем есть ли ФИО в базе ИП и переключаемся на другое ФИО если находим
        name_user = name_act
    else:
        name_user = new_name

    name_user = name_user.replace(" ", "")
    # name_project = input('Введите ключ проекта: \n')
    name_project = ''

    # ----------------------- получение задач проекта-----------------------
    all_project = jira.projects()                   #получение всех проектов
    pa = project_act.replace(' ', '').lower()       #форматирование project_act
    for p in range(len(all_project)):
        n = all_project[p].name.replace(' ', '').lower()
        if pa == n:                                 #если наименование проекта найдено
            name_project = all_project[p].key
    print(name_project)
    #print(date.today() - timedelta(35))
    if name_project != '':                          #если наименование проекта найдено
        issues_in_proj = []
        issues_in_proj = jira.search_issues(f'project={name_project} and updated > {datedelta}', maxResults=300)
        proj = jira.project(name_project)
        project_jira = proj.name
        #print(proj.name)  # имя проекта
        #print(issues_in_proj)
        #print(len(issues_in_proj))

        # ----------------------- поиск задач в которых есть ФИО-----------------------
        total_time_jira = 0
        for j in range(len(issues_in_proj)):
            x = jira.worklogs(issue=issues_in_proj[j])

            time = 0
            author = '-'

            #print(issues_in_proj[j])
            for i in range(len(x)):
                author = x[i].updateAuthor.displayName
                str = author.replace(" ", "")
                str = str.replace("ё", "е")
                index = x[i].started.split('T')  # разделение даты от времени
                date_jira = datetime.strptime(index[0], "%Y-%m-%d")  # дата джиры
                #print(author)
                if str == name_user and date1 <= date_jira and date2 >= date_jira:
                    #print(x[i].started)
                    #print(index)
                    #print(x[i].id)
                    #print(x[i].timeSpentSeconds)
                    #print(x[i].updateAuthor)
                    time = x[i].timeSpentSeconds + time
            #print(time/3600)
            #print(j)

            total_time_jira = total_time_jira + time


        total_time_jira = format(total_time_jira / 3600, '.2f')
        print('Трудозатраты в джире = ', total_time_jira)
        print('OK')
    # ----------------------------------------- проверка данных ------------------------------------
        print('>>> Проверка данных...')
        text_message = []
        text_message = verific(time_act, total_time_jira, project_act, project_jira, date_act_f, full_name_act1, name_act2, name_act3, number_act, number_zayavka, rate_act, rate_zayavka, cost_for_verification, project_cost_act, total_cost_act, total_cost_act_in_text, total_cost_zayavka, date1, date2)
        print('ОК')
    #----------------------------------------- создание отчета------------------------------------
        report.create_report(date_act_f, number_act, number_zayavka, project_act, key_project_act, period, time_act, rate_act, rate_zayavka, project_cost_act, total_cost_act, total_cost_zayavka, total_cost_act_in_text, name_act2, name_act3, name_act, project_jira, date_start, date_end, total_time_jira, cost_for_verification, name_act)
    # ----------------------------------------- отправка сообщения ------------------------------------
        print('>>> Отправка сообщения...')
        send_message(text_message[0], name_act, path, name_act2, number_act, date_for_rename, total_cost_act, text_message[1], type_of_act, project_act)
        print('ОК')
    if name_project == '':  # если наименование проекта не найдено
        send_message('Завершена роботизированная проверка акта и заявки.\nРезультат обработки:'+ '\n\n' +
                     '     Проект, указанный в акте, не найден в JIRA. Проверьте, пожалуйста, верно ли указано наименование проекта в акте.' + '\n\n' +
                     'Акт и заявка во вложении. Просьба проверить документы в соответствии с замечаниями и направить повторно на проверку на электронный адрес actbot@i-sol.ru', '', path, '', '', '', '', 0, '', '')
        print('>>> Проект не найден, сообщение отправлено')

input("Нажмите enter для завершения...")